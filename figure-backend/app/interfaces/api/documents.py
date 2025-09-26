"""
Documents API Endpoints
문서 업로드 및 관리 API
"""
import logging
import uuid
import os
import json
from typing import List, Dict, Any, Optional
from datetime import datetime
from fastapi import APIRouter, HTTPException, Depends, status, UploadFile, File, Form, Request, Query
from fastapi.responses import FileResponse
import aiofiles

from app.domain.entities.schemas import (
    Document, UploadDocumentRequest, APIResponse, DocumentType,
    JobCreate, JobType, JobUpdate, JobStatus
)
from app.domain.entities.template_entities import (
    Template, TemplateType, TemplateFormat, TemplateCreateRequest
)
from app.infrastructure.dependencies.logging_dependencies import business_logger_dependency
from app.infrastructure.middleware.logging_middleware import BusinessEventLogger
from app.application.services.vector_store import vector_store_service
from app.application.services.job_service import job_service
from app.application.services.template_service import get_template_service, TemplateService
from app.config import settings

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/documents", tags=["Documents"])

# 업로드 파일 저장 경로 (절대 경로로 설정)
UPLOAD_DIRECTORY = "/app/data/uploads"
os.makedirs(UPLOAD_DIRECTORY, exist_ok=True)

# 지원하는 파일 형식
SUPPORTED_EXTENSIONS = {
    '.txt': 'text',
    '.md': 'text',
    '.pdf': 'pdf',
    '.doc': 'doc',
    '.docx': 'doc',
    '.html': 'website',
    '.htm': 'website'
}

# 프론트엔드 템플릿 타입 매핑
FRONTEND_TEMPLATE_TYPE_MAPPING = {
    # 1단계: 요구사항 분석
    'REQUIREMENTS': 'requirements',
    'BUSINESS_REQUIREMENTS': 'business-requirements',
    'FUNCTIONAL_SPECIFICATION': 'functional-specification',
    
    # 2단계: 설계
    'TECHNICAL_SPECIFICATION': 'technical-specification',
    'SYSTEM_ARCHITECTURE': 'system-architecture',
    'DATABASE_DESIGN': 'database-design',
    'TABLE_SPECIFICATION': 'table-specification',
    'API_SPECIFICATION': 'api-specification',
    'UI_UX_DESIGN': 'ui-ux-design',
    
    # 3단계: 개발
    'IMPACT_ANALYSIS': 'impact-analysis',
    'API_DOCUMENTATION': 'api-documentation',
    'CODE_REVIEW_CHECKLIST': 'code-review-checklist',
    
    # 4단계: 테스트
    'TEST_PLAN': 'test-plan',
    'TEST_SCENARIO': 'test-scenario',
    'TEST_CASE': 'test-case',
    'QA_CHECKLIST': 'qa-checklist',
    
    # 5단계: 배포
    'DEPLOYMENT_GUIDE': 'deployment-guide',
    'DEPLOYMENT_CHECKLIST': 'deployment-checklist',
    'ROLLBACK_PLAN': 'rollback-plan',
    'MONITORING_PLAN': 'monitoring-plan',
    
    # 6단계: 운영
    'USER_MANUAL': 'user-manual',
    'RELEASE_NOTES': 'release-notes',
    'OPERATION_MANUAL': 'operation-manual',
    'TROUBLESHOOTING_GUIDE': 'troubleshooting-guide',
    
    'CUSTOM': 'custom'
}

async def get_vector_store_service():
    """벡터 스토어 서비스 의존성 주입"""
    if not vector_store_service._collection:
        await vector_store_service.initialize()
    return vector_store_service


@router.post(
    "/upload-file",
    response_model=APIResponse,
    summary="파일 업로드 (하이브리드) - 백오피스 호환",
    description="백오피스에서 보내는 형식에 맞춰 파일을 업로드하고 벡터 데이터베이스에 저장합니다. 템플릿인 경우 SQLite에도 저장됩니다."
)
async def upload_file(
    request: Request,
    file: UploadFile = File(..., description="업로드할 파일"),
    site_id: Optional[str] = Form(default=None, description="관련 사이트 ID"),
    metadata: Optional[str] = Form(default="{}", description="메타데이터 JSON 문자열 (템플릿 정보 포함)"),
    service: object = Depends(get_vector_store_service),
    template_service: TemplateService = Depends(get_template_service),
    business_logger: BusinessEventLogger = Depends(business_logger_dependency)
) -> APIResponse:
    """
    백오피스 호환 파일 업로드
    
    **프론트엔드에서 보내는 데이터 구조:**
    - **file**: 업로드할 파일 (필수)
    - **site_id**: 관련 사이트 ID (필수)
    - **metadata**: JSON 문자열로 다음 정보 포함:
      - description: 파일 설명
      - tags: 태그 배열
      - template_type: 템플릿 유형 (REQUIREMENTS, IMPACT_ANALYSIS 등 - 필수)
      - template_version: 템플릿 버전 (기본값: "1.0.0")
    """
    # 로그 컨텍스트 가져오기
    log_context = getattr(request.state, 'log_context', None)
    
    # 디버깅을 위한 요청 데이터 로깅
    logger.info(f"파일 업로드 요청 받음")
    logger.info(f"  - 파일명: {file.filename if file else 'None'}")
    logger.info(f"  - site_id 원본값: '{site_id}'")
    logger.info(f"  - site_id 타입: {type(site_id)}")
    logger.info(f"  - metadata 원본값: '{metadata}'")
    
    try:
        # site_id 검증 (필수 항목) - 빈 문자열도 체크
        if not site_id or site_id == "" or site_id.strip() == "":
            logger.error(f"site_id 검증 실패: '{site_id}'")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="사이트 ID는 필수 항목입니다."
            )
        
        # 파일 기본 검증
        if not file or not file.filename:
            logger.error(f"파일 검증 실패: file={file}, filename={file.filename if file else 'None'}")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="파일명이 제공되지 않았습니다."
            )
        
        # 파일 확장자 확인
        file_ext = os.path.splitext(file.filename)[1].lower() 
        if file_ext not in SUPPORTED_EXTENSIONS:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"지원하지 않는 파일 형식입니다. 지원 형식: {', '.join(SUPPORTED_EXTENSIONS.keys())}"
            )
        
        # 메타데이터 파싱
        try:
            # metadata가 None이거나 빈 문자열인 경우 처리
            if not metadata or metadata == "":
                parsed_metadata = {}
            else:
                parsed_metadata = json.loads(metadata) if metadata != "{}" else {}
        except (json.JSONDecodeError, TypeError) as e:
            logger.error(f"메타데이터 파싱 오류: {e}, metadata: {metadata}")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"메타데이터 JSON 형식이 올바르지 않습니다: {str(e)}"
            )
        
        # 메타데이터에서 템플릿 정보 추출
        template_type_frontend = parsed_metadata.get('template_type', '')
        template_version = parsed_metadata.get('template_version', '1.0.0')
        description = parsed_metadata.get('description', '')
        tags = parsed_metadata.get('tags', [])
        
        # 템플릿 여부 판단
        is_template = bool(template_type_frontend)
        
        # template_type 검증 (필수)
        if not template_type_frontend or template_type_frontend == "":
            logger.error(f"template_type 검증 실패: '{template_type_frontend}'")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="템플릿 유형은 필수 항목입니다."
            )
        

            
        # 프론트엔드 템플릿 타입을 백엔드 타입으로 매핑
        template_type_backend = FRONTEND_TEMPLATE_TYPE_MAPPING.get(
            template_type_frontend, 
            template_type_frontend.lower().replace('_', '-') if template_type_frontend else ''
        )
        
        # 템플릿 이름 생성 (파일명 기반)
        template_name = os.path.splitext(file.filename)[0]
        
        # 비즈니스 이벤트 로깅: 파일 업로드 시작
        if log_context:
            await business_logger.logging_service.log_business_event(
                log_context,
                "파일_업로드_시작",
                {
                    "filename": file.filename,
                    "content_type": file.content_type,
                    "site_id": site_id,
                    "is_template": is_template,
                    "template_type": template_type_backend,
                    "template_name": template_name
                }
            )

        
        # 파일 크기 확인 (10MB 제한)
        MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
        content = await file.read()
        if len(content) > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                detail="파일 크기가 10MB를 초과합니다."
            )
        
        # 작업 추적 시작
        job = job_service.create_job(JobCreate(
            type=JobType.DOCUMENT_UPLOAD,
            site_id=site_id,
            metadata={
                "filename": file.filename,
                "file_size": len(content),
                "file_type": file_ext
            }
        ))
        
        try:
            # 진행 상태 업데이트: 파일 저장 중
            job_service.update_job(job.id, JobUpdate(
                status=JobStatus.PROCESSING,
                progress=20.0,
                message="파일 저장 중..."
            ))
            
            # 파일 저장
            file_id = str(uuid.uuid4())
            file_path = os.path.join(UPLOAD_DIRECTORY, f"{file_id}{file_ext}")
            
            async with aiofiles.open(file_path, 'wb') as f:
                await f.write(content)
            
            # 진행 상태 업데이트: 내용 추출 중
            job_service.update_job(job.id, JobUpdate(
                progress=40.0,
                message="내용 추출 중..."
            ))
            
            # 파일 내용 추출
            file_content = await extract_file_content(file_path, file_ext)
            
            # 진행 상태 업데이트: 문서 생성 중
            job_service.update_job(job.id, JobUpdate(
                progress=60.0,
                message="문서 생성 중..."
            ))
            
                # ChromaDB 호환을 위해 배열 타입을 문자열로 변환
            processed_metadata = {}
            for key, value in parsed_metadata.items():
                if isinstance(value, list):
                    processed_metadata[key] = ", ".join(str(item) for item in value)
                elif isinstance(value, dict):
                    processed_metadata[key] = json.dumps(value)
                else:
                    processed_metadata[key] = value
            
            # Document 객체 생성 (모든 메타데이터 포함)
            document = Document(
                id=file_id,
                title=file.filename,
                content=file_content,
                doc_type=DocumentType(SUPPORTED_EXTENSIONS[file_ext]),
                site_id=site_id,
                metadata={
                    "file_path": file_path,
                    "file_size": len(content),
                    "original_filename": file.filename,
                    "template_type": template_type_backend,
                    "template_name": template_name,
                    "template_description": description,
                    "template_version": template_version,
                    "is_template": is_template,
                    "description": description,
                    "tags": ", ".join(tags) if isinstance(tags, list) else tags,
                    **processed_metadata
                },
                source_url=file_path,
                created_at=datetime.now()
            )
            
            # 진행 상태 업데이트: 벡터화 중
            job_service.update_job(job.id, JobUpdate(
                progress=80.0,
                message="벡터화 처리 중..."
            ))
            
            # 벡터 저장소에 추가
            doc_id = await service.add_document(document)
            
            # 템플릿인 경우 SQLite에도 저장 (하이브리드 저장)
            template_id = None
            if is_template and template_type_backend:
                try:
                    # 진행 상태 업데이트: 템플릿 저장 중
                    job_service.update_job(job.id, JobUpdate(
                        progress=90.0,
                        message="템플릿 저장 중..."
                    ))
                    
                    # 기존 템플릿 버전 확인 및 자동 증가
                    base_name = template_name
                    existing_versions = []
                    
                    try:
                        # 백엔드 템플릿 타입으로 변환
                        try:
                            template_type_enum = TemplateType(template_type_backend)
                        except ValueError:
                            # 매핑되지 않은 타입은 CUSTOM으로 처리
                            template_type_enum = TemplateType.CUSTOM
                        
                        # 기존 템플릿들 조회
                        existing_templates = await template_service.get_templates_by_filters(
                            template_type=template_type_enum,
                            site_id=site_id
                        )
                        
                        # 같은 기본 이름을 가진 템플릿들의 버전 수집
                        for template in existing_templates:
                            if template.name.startswith(base_name):
                                existing_versions.append(template.version)
                                
                    except Exception as e:
                        logger.warning(f"기존 템플릿 버전 확인 실패: {e}")
                    
                    # 버전 자동 증가 로직
                    final_version = template_version
                    if existing_versions:
                        # 기존 버전이 있으면 자동으로 다음 버전 생성
                        try:
                            # 시맨틱 버전 파싱 및 증가
                            latest_version = max(existing_versions, key=lambda v: [int(x) for x in v.split('.')])
                            version_parts = latest_version.split('.')
                            patch_version = int(version_parts[2]) + 1 if len(version_parts) >= 3 else 1
                            final_version = f"{version_parts[0]}.{version_parts[1]}.{patch_version}"
                            
                            logger.info(f"자동 버전 증가: {latest_version} -> {final_version}")
                        except Exception as e:
                            logger.warning(f"버전 자동 증가 실패, 기본 버전 사용: {e}")
                    
                    # 템플릿 생성 요청 구성
                    template_request = TemplateCreateRequest(
                        name=f"{base_name}_v{final_version}",
                        description=description or f"{file.filename}에서 추출한 템플릿 (v{final_version})",
                        template_type=template_type_enum,
                        format=TemplateFormat.MARKDOWN if file_ext in ['.md'] else TemplateFormat.TEXT,
                        site_id=site_id,
                        content=file_content,
                        variables={},
                        tags=tags + ["uploaded", "hybrid", f"v{final_version}"] if isinstance(tags, list) else ["uploaded", "hybrid", f"v{final_version}"],
                        is_default=False,
                        version=final_version
                    )
                    
                    # 템플릿 저장
                    template = await template_service.create_template(
                        request=template_request,
                        created_by="hybrid_upload",
                        file_content=content,
                        file_name=file.filename,
                        content_type=file.content_type
                    )
                    template_id = template.id
                    
                    logger.info(f"하이브리드 저장 완료 - 문서: {doc_id}, 템플릿: {template_id}")
                    
                except Exception as template_error:
                    logger.warning(f"템플릿 저장 실패 (벡터 저장은 성공): {template_error}")
                    # 템플릿 저장 실패해도 벡터 저장은 성공했으므로 계속 진행
            
            # 진행 상태 업데이트: 완료
            job_service.update_job(job.id, JobUpdate(
                status=JobStatus.COMPLETED,
                progress=100.0,
                message="파일 업로드 완료" + (" (하이브리드 저장)" if is_template else ""),
                metadata={
                    "document_id": doc_id,
                    "template_id": template_id
                }
            ))
            
            logger.info(f"파일 업로드 성공: {file.filename} -> {doc_id}" + 
                       (f", 템플릿: {template_id}" if template_id else ""))
            
            # 응답 데이터 구성 (프론트엔드 호환)
            response_data = {
                "document_id": doc_id,
                "job_id": job.id,
                "filename": file.filename,
                "file_size": len(content),
                "doc_type": SUPPORTED_EXTENSIONS[file_ext],
                "created_at": document.created_at.isoformat(),
                "is_hybrid": is_template,
                "is_template": is_template,
                "template_type": template_type_backend if is_template else None,
                "template_name": template_name if is_template else None,
                "template_version": final_version if is_template and 'final_version' in locals() else template_version,
                "site_id": site_id,
                "description": description,
                "tags": tags
            }
            
            if template_id:
                response_data["template_id"] = template_id
            
            return APIResponse(
                success=True,
                message="파일 업로드 성공" + (" (하이브리드 저장 완료)" if is_template else ""),
                data=response_data
            )
            
        except Exception as e:
            # 진행 상태 업데이트: 실패
            job_service.update_job(job.id, JobUpdate(
                status=JobStatus.FAILED,
                error=str(e),
                message="파일 업로드 실패"
            ))
            
            # 실패 시 파일 삭제
            if 'file_path' in locals() and os.path.exists(file_path):
                os.remove(file_path)
            
            raise
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"파일 업로드 실패: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"파일 업로드 실패: {str(e)}"
        )


async def extract_file_content(file_path: str, file_ext: str) -> str:
    """파일 내용 추출"""
    try:
        if file_ext in ['.txt', '.md', '.html', '.htm']:
            # 텍스트 파일
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                return await f.read()
        elif file_ext == '.pdf':
            # PDF 파일 처리
            try:
                import PyPDF2
                content = ""
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        content += page.extract_text() + "\n"
                return content.strip() if content.strip() else f"PDF 파일에서 텍스트를 추출할 수 없습니다: {os.path.basename(file_path)}"
            except ImportError:
                logger.error("PyPDF2 라이브러리가 설치되지 않았습니다")
                return f"PDF 처리 라이브러리가 설치되지 않았습니다: {os.path.basename(file_path)}"
            except Exception as e:
                logger.error(f"PDF 파일 처리 중 오류: {e}")
                return f"PDF 파일 처리 중 오류가 발생했습니다: {str(e)}"
        elif file_ext in ['.doc', '.docx']:
            # Word 파일 처리
            try:
                from docx import Document
                doc = Document(file_path)
                content = ""
                for paragraph in doc.paragraphs:
                    content += paragraph.text + "\n"
                
                # 표 내용도 추출
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            content += cell.text + " "
                        content += "\n"
                
                return content.strip() if content.strip() else f"Word 파일에서 텍스트를 추출할 수 없습니다: {os.path.basename(file_path)}"
            except ImportError:
                logger.error("python-docx 라이브러리가 설치되지 않았습니다")
                return f"Word 처리 라이브러리가 설치되지 않았습니다: {os.path.basename(file_path)}"
            except Exception as e:
                logger.error(f"Word 파일 처리 중 오류: {e}")
                return f"Word 파일 처리 중 오류가 발생했습니다: {str(e)}"
        else:
            return f"지원하지 않는 파일 형식: {file_ext}"
    except Exception as e:
        logger.error(f"파일 내용 추출 실패: {e}")
        return f"파일 내용 추출 실패: {str(e)}"


@router.get(
    "/download/{document_id}",
    summary="파일 다운로드",
    description="업로드된 파일을 다운로드합니다."
)
async def download_file(
    document_id: str,
    service: object = Depends(get_vector_store_service)
):
    """파일 다운로드"""
    try:
        # 문서 정보 조회 (실제 구현에서는 데이터베이스에서 조회)
        file_path = os.path.join(UPLOAD_DIRECTORY, f"{document_id}.txt")  # 임시로 txt 확장자
        
        if not os.path.exists(file_path):
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="파일을 찾을 수 없습니다."
            )
        
        return FileResponse(
            file_path,
            media_type='application/octet-stream',
            filename=f"{document_id}.txt"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"파일 다운로드 실패: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"파일 다운로드 실패: {str(e)}"
        )


@router.post(
    "/upload",
    response_model=APIResponse,
    summary="문서 업로드 (하이브리드)",
    description="새 문서를 벡터 데이터베이스에 업로드하고, 템플릿인 경우 SQLite에도 저장합니다."
)
async def upload_document(
    request: UploadDocumentRequest,
    service: object = Depends(get_vector_store_service),
    template_service: TemplateService = Depends(get_template_service)
) -> APIResponse:
    """
    문서 업로드
    
    - **title**: 문서 제목
    - **content**: 문서 내용 (최소 10자)
    - **doc_type**: 문서 타입 (website, pdf, text, confluence, jira)
    - **source_url**: 소스 URL (선택사항)
    - **site_id**: 관련 사이트 ID (선택사항)
    - **metadata**: 추가 메타데이터 (선택사항)
    """
    try:
        # 작업 추적 시작
        job = job_service.create_job(JobCreate(
            type=JobType.DOCUMENT_UPLOAD,
            site_id=request.site_id,
            metadata={
                "title": request.title,
                "doc_type": request.doc_type.value,
                "content_length": len(request.content)
            }
        ))
        
        try:
            # 진행 상태 업데이트: 시작
            job_service.update_job(job.id, JobUpdate(
                status=JobStatus.PROCESSING,
                progress=10.0,
                message="문서 검증 중..."
            ))
            
            # Document 객체 생성
            document = Document(
                id=str(uuid.uuid4()),
                title=request.title,
                content=request.content,
                doc_type=request.doc_type,
                source_url=request.source_url,
                site_id=request.site_id,
                metadata=request.metadata,
                created_at=datetime.now()
            )
            
            # 진행 상태 업데이트: 벡터화 중
            job_service.update_job(job.id, JobUpdate(
                progress=50.0,
                message="벡터화 처리 중..."
            ))
            
            # 벡터 저장소에 추가
            doc_id = await service.add_document(document)
            
            # 통합 DB 저장: 문서를 figure.db에 저장 (템플릿 정보 포함)
            template_id = None
            logger.info(f"통합 DB 저장 체크 - metadata: {request.metadata}, type: {type(request.metadata)}")
            
            # 문서를 figure.db에 저장 (템플릿 정보 포함)
            try:
                from app.infrastructure.persistence.connection import DatabaseManager
                from app.infrastructure.persistence.models import Document as DocumentModel
                
                # 메타데이터에서 템플릿 정보 추출
                is_template = False
                template_type_str = ""
                template_version = "1.0.0"
                template_description = ""
                template_tags = ""
                
                if request.metadata and isinstance(request.metadata, dict):
                    template_type_str = request.metadata.get('template_type', '')
                    is_template = bool(template_type_str)
                    template_version = request.metadata.get('template_version', '1.0.0')
                    template_description = request.metadata.get('description', request.title or '')
                    tags = request.metadata.get('tags', '')
                    
                    # 태그 처리 (문자열을 리스트로 변환)
                    if isinstance(tags, str):
                        template_tags = tags
                    elif isinstance(tags, list):
                        template_tags = ",".join(tags)
                    else:
                        template_tags = ""
                
                logger.info(f"템플릿 체크 - template_type: '{template_type_str}', is_template: {is_template}")
                
                # figure.db에 문서 저장
                db_manager = DatabaseManager()
                with db_manager.SessionLocal() as db_session:
                    # 템플릿 타입 매핑
                    template_type_backend = FRONTEND_TEMPLATE_TYPE_MAPPING.get(
                        template_type_str, 
                        template_type_str.lower().replace('_', '-') if template_type_str else ''
                    )
                    
                    # Document 모델 생성
                    db_document = DocumentModel(
                        id=doc_id,
                        site_id=request.site_id,
                        title=request.title,
                        doc_type=request.doc_type.value,
                        source_url=request.source_url or "",
                        file_path=None,  # JSON 업로드는 파일 경로 없음
                        file_size=len(request.content),
                        content_hash=None,  # 필요시 나중에 추가
                        processing_status="completed",
                        extra_data=request.metadata or {},
                        
                        # 템플릿 관련 필드
                        is_template=is_template,
                        template_type=template_type_backend if is_template else None,
                        template_name=request.title if is_template else None,
                        template_description=template_description if is_template else None,
                        template_version=template_version if is_template else "1.0.0",
                        template_format="markdown" if request.doc_type.value == 'text' else "text",
                        template_variables={},
                        template_tags=template_tags if is_template else None,
                        usage_count=0,
                        is_active=True,
                        is_default=False,
                        created_by="hybrid_upload_json"
                    )
                    
                    # DB에 저장
                    db_session.add(db_document)
                    db_session.commit()
                    
                    logger.info(f"통합 DB 저장 완료 - 문서: {doc_id}, 템플릿: {is_template}")
                    
                    if is_template:
                        template_id = doc_id  # 문서 ID를 템플릿 ID로 사용
                        
            except Exception as db_error:
                logger.warning(f"통합 DB 저장 실패 (벡터 저장은 성공): {db_error}")
                

            
            # 진행 상태 업데이트: 완료
            job_service.update_job(job.id, JobUpdate(
                status=JobStatus.COMPLETED,
                progress=100.0,
                message="문서 업로드 완료" + (f" (템플릿 ID: {template_id})" if template_id else ""),
                metadata={"document_id": doc_id, "template_id": template_id}
            ))
            
            logger.info(f"문서 업로드 성공: {doc_id}")
            
            return APIResponse(
                success=True,
                message="문서 업로드 성공" + (" (하이브리드 저장)" if template_id else ""),
                data={
                    "document_id": doc_id,
                    "template_id": template_id,
                    "job_id": job.id,
                    "title": document.title,
                    "doc_type": document.doc_type.value,
                    "created_at": document.created_at.isoformat(),
                    "is_hybrid": template_id is not None
                }
            )
            
        except Exception as e:
            # 진행 상태 업데이트: 실패
            job_service.update_job(job.id, JobUpdate(
                status=JobStatus.FAILED,
                error=str(e),
                message="문서 업로드 실패"
            ))
            raise
        
    except ValueError as e:
        logger.error(f"문서 업로드 검증 실패: {e}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(e)
        )
    except Exception as e:
        logger.error(f"문서 업로드 실패: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"문서 업로드 실패: {str(e)}"
        )


@router.delete(
    "/{document_id}",
    response_model=APIResponse,
    summary="문서 삭제",
    description="지정된 문서를 벡터 데이터베이스에서 삭제합니다."
)
async def delete_document(
    document_id: str,
    service: object = Depends(get_vector_store_service)
) -> APIResponse:
    """문서 삭제"""
    try:
        success = await service.delete_document(document_id)
        
        if success:
            logger.info(f"문서 삭제 성공: {document_id}")
            return APIResponse(
                success=True,
                message="문서 삭제 성공",
                data={"document_id": document_id}
            )
        else:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"문서를 찾을 수 없습니다: {document_id}"
            )
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"문서 삭제 실패: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"문서 삭제 실패: {str(e)}"
        )


@router.get(
    "",
    response_model=APIResponse,
    summary="문서 목록 조회",
    description="벡터 데이터베이스에 저장된 모든 문서 목록을 조회합니다."
)
async def get_documents(
    service: object = Depends(get_vector_store_service)
) -> APIResponse:
    """문서 목록 조회"""
    try:
        # ChromaDB에서 모든 문서 정보 가져오기
        collection_info = await service.get_collection_info()
        
        # 컬렉션이 비어있다면 빈 목록 반환
        if collection_info.get("total_chunks", 0) == 0:
            return APIResponse(
                success=True,
                message="문서 목록 조회 성공",
                data={"documents": [], "total": 0}
            )
        
        # 모든 문서 조회 (ChromaDB의 get() 메서드 사용)
        collection = service._collection
        if collection is None:
            return APIResponse(
                success=True,
                message="문서 목록 조회 성공",
                data={"documents": [], "total": 0}
            )
        
        # 모든 문서 ID와 메타데이터 가져오기
        results = collection.get()
        
        # 디버깅 로그 추가
        logger.info(f"ChromaDB 조회 결과: {len(results.get('ids', []))}개의 ID")
        logger.info(f"IDs 샘플: {results.get('ids', [])[:3]}")
        
        # 문서 정보 구성
        documents = []
        processed_doc_ids = set()
        
        for i, doc_id in enumerate(results.get("ids", [])):
            metadata = results.get("metadatas", [{}])[i] if i < len(results.get("metadatas", [])) else {}
            
            # 원본 문서 ID 추출 (chunk ID에서 _chunk_X 부분 제거)
            original_doc_id = doc_id.split("_chunk_")[0] if "_chunk_" in doc_id else doc_id
            
            # 이미 처리된 문서는 건너뛰기 (중복 제거)
            if original_doc_id in processed_doc_ids:
                continue
                
            processed_doc_ids.add(original_doc_id)
            
            document_info = {
                "id": original_doc_id,
                "filename": metadata.get("title", "제목 없음"),  # 프론트엔드가 기대하는 필드명
                "type": metadata.get("doc_type", "unknown"),     # 프론트엔드가 기대하는 필드명
                "created_at": metadata.get("created_at"),
                "updated_at": metadata.get("created_at"),        # 현재는 생성일과 동일
                "site_id": metadata.get("site_id"),
                "source_url": metadata.get("source_url"),
                "vector_count": metadata.get("total_chunks", 1), # 프론트엔드가 기대하는 필드명
                "size": metadata.get("file_size", 0),           # 파일 크기 (없으면 0)
                "status": "processed",                           # 일단 모든 문서는 processed 상태로
                # 백워드 호환성을 위해 기존 필드도 유지
                "title": metadata.get("title", "제목 없음"),
                "doc_type": metadata.get("doc_type", "unknown"),
                "total_chunks": metadata.get("total_chunks", 1)
            }
            
            documents.append(document_info)
        
        # 생성일시 기준 내림차순 정렬
        documents.sort(key=lambda x: x.get("created_at", ""), reverse=True)
        
        return APIResponse(
            success=True,
            message="문서 목록 조회 성공",
            data={
                "documents": documents,
                "total": len(documents)
            }
        )
        
    except Exception as e:
        logger.error(f"문서 목록 조회 실패: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"문서 목록 조회 실패: {str(e)}"
        )


@router.get(
    "/list",
    response_model=APIResponse,
    summary="문서 목록 조회",
    description="벡터 데이터베이스에 저장된 모든 문서 목록을 조회합니다."
)
async def list_documents(
    service: object = Depends(get_vector_store_service)
) -> APIResponse:
    """문서 목록 조회"""
    try:
        # ChromaDB에서 모든 문서 정보 가져오기
        collection_info = await service.get_collection_info()
        
        # 컬렉션이 비어있다면 빈 목록 반환
        if collection_info.get("total_documents", 0) == 0:
            return APIResponse(
                success=True,
                message="문서 목록 조회 성공",
                data={"documents": [], "total": 0}
            )
        
        # 모든 문서 조회 (ChromaDB의 get() 메서드 사용)
        collection = service._collection
        if collection is None:
            return APIResponse(
                success=True,
                message="문서 목록 조회 성공",
                data={"documents": [], "total": 0}
            )
        
        # 모든 문서 ID와 메타데이터 가져오기
        results = collection.get()

                # 디버깅 로그 추가
        logger.info(f"ChromaDB 조회 결과: {len(results.get('ids', []))}개의 ID")
        logger.info(f"IDs 샘플: {results.get('ids', [])[:3]}")
        
        # 문서 정보 구성
        documents = []
        processed_doc_ids = set()
        
        for i, doc_id in enumerate(results.get("ids", [])):
            metadata = results.get("metadatas", [{}])[i] if i < len(results.get("metadatas", [])) else {}
            
            # 원본 문서 ID 추출 (chunk ID에서 _chunk_X 부분 제거)
            original_doc_id = doc_id.split("_chunk_")[0] if "_chunk_" in doc_id else doc_id
            
            # 이미 처리된 문서는 건너뛰기 (중복 제거)
            if original_doc_id in processed_doc_ids:
                continue
                
            processed_doc_ids.add(original_doc_id)
            
            document_info = {
                "id": original_doc_id,
                "filename": metadata.get("title", "제목 없음"),  # 프론트엔드가 기대하는 필드명
                "type": metadata.get("doc_type", "unknown"),     # 프론트엔드가 기대하는 필드명
                "created_at": metadata.get("created_at"),
                "updated_at": metadata.get("created_at"),        # 현재는 생성일과 동일
                "site_id": metadata.get("site_id"),
                "source_url": metadata.get("source_url"),
                "vector_count": metadata.get("total_chunks", 1), # 프론트엔드가 기대하는 필드명
                "size": metadata.get("file_size", 0),           # 파일 크기 (없으면 0)
                "status": "processed",                           # 일단 모든 문서는 processed 상태로
                # 백워드 호환성을 위해 기존 필드도 유지
                "title": metadata.get("title", "제목 없음"),
                "doc_type": metadata.get("doc_type", "unknown"),
                "total_chunks": metadata.get("total_chunks", 1)
            }
            
            documents.append(document_info)
        
        # 생성일시 기준 내림차순 정렬
        documents.sort(key=lambda x: x.get("created_at", ""), reverse=True)
        
        return APIResponse(
            success=True,
            message="문서 목록 조회 성공",
            data={
                "documents": documents,
                "total": len(documents)
            }
        )
        
    except Exception as e:
        logger.error(f"문서 목록 조회 실패: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"문서 목록 조회 실패: {str(e)}"
        )


@router.get(
    "/search",
    response_model=APIResponse,
    summary="문서 검색",
    description="유사도 기반으로 관련 문서를 검색합니다."
)
async def search_documents(
    query: str,
    max_results: int = 5,
    similarity_threshold: float = 0.7,
    site_ids: str = None,  # 쿼리 파라미터로 쉼표 구분 문자열
    service: object = Depends(get_vector_store_service)
) -> APIResponse:
    """
    문서 검색
    
    - **query**: 검색어
    - **max_results**: 최대 결과 수 (기본값: 5)
    - **similarity_threshold**: 유사도 임계값 (기본값: 0.7)
    - **site_ids**: 사이트 ID 목록 (쉼표로 구분, 선택사항)
    """
    try:
        # site_ids 파싱
        site_id_list = None
        if site_ids:
            site_id_list = [s.strip() for s in site_ids.split(",") if s.strip()]
        
        # 검색 실행
        results = await service.search_similar(
            query=query,
            max_results=max_results,
            site_ids=site_id_list,
            similarity_threshold=similarity_threshold
        )
        
        return APIResponse(
            success=True,
            message="문서 검색 성공",
            data={
                "query": query,
                "results": results,
                "total_found": len(results)
            }
        )
        
    except Exception as e:
        logger.error(f"문서 검색 실패: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"문서 검색 실패: {str(e)}"
        )


@router.get(
    "/search-raw", 
    response_model=APIResponse,
    summary="문서 원본 검색 (MCP용)",
    description="벡터 검색만 수행하고 LLM 처리 없이 원본 문서들을 반환합니다. Cursor/Claude가 직접 처리할 수 있습니다."
)
async def search_documents_raw(
    query: str = Query(..., description="검색 질의"),
    max_results: int = Query(5, description="최대 결과 수", le=20),
    similarity_threshold: float = Query(0.3, description="유사도 임계값 (0.0-1.0)", ge=0.0, le=1.0),
    site_ids: str = Query(None, description="사이트 ID 목록 (쉼표로 구분)"),
    service: object = Depends(get_vector_store_service)
) -> APIResponse:
    """
    MCP 전용 원본 문서 검색
    
    - Cursor/Claude가 직접 RAG 처리할 수 있도록 원본 문서만 반환
    - LLM 처리 없이 벡터 검색 결과만 제공
    - 빠른 응답 시간과 효율적인 처리
    """
    try:
        # site_ids 파싱
        site_id_list = None
        if site_ids:
            site_id_list = [s.strip() for s in site_ids.split(",") if s.strip()]

        logger.info(f"원본 문서 검색 요청: query='{query}', max_results={max_results}, threshold={similarity_threshold}")
        
        # 벡터 검색만 수행 (LLM 처리 없음)
        search_results = await service.search_similar(
            query=query,
            max_results=max_results,
            site_ids=site_id_list,
            similarity_threshold=similarity_threshold
        )
        
        if not search_results:
            return APIResponse(
                success=True,
                message=f"'{query}'와 관련된 문서를 찾을 수 없습니다",
                data=[]
            )
        
        # MCP가 처리할 수 있는 형태로 데이터 구성
        formatted_results = []
        for result in search_results:
            formatted_result = {
                "content": result["content"],
                "similarity": result["similarity"],
                "rank": result["rank"]
            }
            
            # 메타데이터 추가 (필요한 것만)
            if "metadata" in result and result["metadata"]:
                metadata = result["metadata"]
                formatted_result["metadata"] = {
                    "title": metadata.get("title"),
                    "site_id": metadata.get("site_id"), 
                    "doc_type": metadata.get("doc_type"),
                    "tags": metadata.get("tags"),
                    "template_type": metadata.get("template_type")
                }
            
            formatted_results.append(formatted_result)
        
        logger.info(f"원본 문서 검색 완료: {len(formatted_results)}개 결과 반환")
        
        return APIResponse(
            success=True,
            message=f"'{query}' 검색 완료: {len(formatted_results)}개 문서 발견",
            data=formatted_results
        )
        
    except Exception as e:
        logger.error(f"원본 문서 검색 실패: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="원본 문서 검색 중 오류가 발생했습니다"
        )


@router.get(
    "/",
    response_model=APIResponse,
    summary="문서 목록 조회",
    description="하이브리드 저장 시스템의 문서 목록을 조회합니다. RDB와 Vector DB 저장 상태를 포함합니다."
)
async def get_documents(
    limit: int = 50,
    offset: int = 0,
    site_id: str = None,
    template_type: str = None,
    service: object = Depends(get_vector_store_service)
) -> APIResponse:
    """
    문서 목록 조회 (하이브리드 저장 상태 포함)
    
    - **limit**: 조회할 문서 수 (기본값: 50)
    - **offset**: 시작 위치 (기본값: 0) 
    - **site_id**: 사이트 ID로 필터링 (선택사항)
    - **template_type**: 템플릿 유형으로 필터링 (선택사항)
    """
    try:
        # Vector DB에서 문서 정보 조회
        collection_info = await service.get_collection_info()
        
        # 필터 조건 구성
        where_conditions = {}
        if site_id:
            where_conditions["site_id"] = site_id
        if template_type:
            where_conditions["template_type"] = template_type
        
        # Vector DB의 모든 문서 메타데이터 조회
        all_results = service._collection.get(
            limit=limit + offset,
            include=["metadatas", "documents"],
            where=where_conditions if where_conditions else None
        )
        
        # 문서별로 그룹화 (chunk들을 문서 단위로 집계)
        documents_map = {}
        
        if all_results["metadatas"]:
            for i, metadata in enumerate(all_results["metadatas"]):
                doc_id = metadata.get("document_id")
                if not doc_id:
                    continue
                

                
                if doc_id not in documents_map:
                    # 템플릿 정보 디버깅
                    template_type = metadata.get("template_type")
                    template_version = metadata.get("template_version")

                    
                    documents_map[doc_id] = {
                        "id": doc_id,
                        "title": metadata.get("title", "제목 없음"),
                        "filename": metadata.get("original_filename", metadata.get("title", "파일명 없음")),
                        "type": metadata.get("doc_type", "unknown"),
                        "site_id": metadata.get("site_id"),
                        "created_at": metadata.get("created_at"),
                        "file_size": metadata.get("file_size", 0),
                        "chunk_count": 0,
                        "total_chunks": metadata.get("total_chunks", 0),
                        # 템플릿 관련 정보
                        "template_type": template_type,
                        "template_version": template_version,
                        "description": metadata.get("description"),
                        "tags": metadata.get("tags"),
                        # 하이브리드 저장 상태
                        "storage_status": {
                            "vector_db": True,  # Vector DB에 있으므로 True
                            "rdb": False,  # 기본값, 아래에서 확인
                            "file_storage": False  # 기본값, 아래에서 확인
                        }
                    }
                else:
                    # 기존 문서에 템플릿 정보가 없거나 빈 문자열이면 업데이트
                    existing_doc = documents_map[doc_id]
                    current_template_type = existing_doc.get("template_type")
                    new_template_type = metadata.get("template_type")
                    
                    # 현재 템플릿 타입이 없거나 빈 문자열이고, 새로운 템플릿 타입이 있으면 업데이트
                    if (not current_template_type or current_template_type == "") and new_template_type:

                        existing_doc["template_type"] = new_template_type
                        existing_doc["template_version"] = metadata.get("template_version")
                        existing_doc["description"] = metadata.get("description") or existing_doc.get("description")
                        existing_doc["tags"] = metadata.get("tags") or existing_doc.get("tags")
                
                documents_map[doc_id]["chunk_count"] += 1
        
        # 문서 목록 생성
        documents_list = list(documents_map.values())[offset:offset+limit]
        
        # 모든 문서에 템플릿 관련 필드가 포함되도록 보장
        for doc in documents_list:
            # 템플릿 관련 필드가 없으면 기본값 설정
            if 'template_type' not in doc or doc.get('template_type') is None:
                doc['template_type'] = ""
            if 'template_version' not in doc or doc.get('template_version') is None:
                doc['template_version'] = ""
            if 'description' not in doc or doc.get('description') is None:
                doc['description'] = ""
            if 'tags' not in doc or doc.get('tags') is None:
                doc['tags'] = ""
                

        
        # 각 문서의 RDB 및 파일 저장 상태 확인
        template_service_instance = get_template_service()
        
        for doc in documents_list:
            # 파일 저장 상태 확인 (파일 시스템)
            file_path = None
            file_exists = False
            
            # 다양한 위치에서 파일 경로 확인
            possible_paths = [
                doc.get("file_path"),
                doc.get("source_url") if doc.get("source_url", "").startswith("/") else None,
                # 업로드 디렉토리에서 문서 ID로 파일 찾기
                f"{UPLOAD_DIRECTORY}/{doc.get('id')}.md",
                f"{UPLOAD_DIRECTORY}/{doc.get('id')}.txt",
                f"{UPLOAD_DIRECTORY}/{doc.get('id')}.pdf",
            ]
            
            # 메타데이터에서도 파일 경로 확인
            if hasattr(doc, '__iter__') and 'metadata' in doc:
                metadata = doc.get("metadata", {})
                if isinstance(metadata, dict):
                    possible_paths.extend([
                        metadata.get("file_path"),
                        metadata.get("original_file_path"),
                        metadata.get("source_url") if metadata.get("source_url", "").startswith("/") else None
                    ])
            
            # 실제 파일 존재 확인
            for path in possible_paths:
                if path and isinstance(path, str):
                    try:
                        if os.path.exists(path):
                            file_path = path
                            file_exists = True
                            break
                    except (OSError, TypeError):
                        continue
            
            # 업로드 디렉토리에서 패턴 매칭으로 파일 찾기 (마지막 시도)
            if not file_exists and doc.get('id'):
                try:
                    doc_id = doc.get('id')
                    for ext in ['.md', '.txt', '.pdf', '.doc', '.docx', '.html']:
                        pattern_path = f"{UPLOAD_DIRECTORY}/{doc_id}{ext}"
                        if os.path.exists(pattern_path):
                            file_path = pattern_path
                            file_exists = True
                            break
                except (OSError, TypeError):
                    pass
            
            doc["storage_status"]["file_storage"] = file_exists
            doc["storage_status"]["file_path"] = file_path if file_exists else None
            
            # RDB 저장 상태 확인 (실제 SQLite 템플릿 테이블에서 확인)
            template_id = doc.get("template_id")
            is_template = doc.get("is_template", False) or doc.get("template_type", "") != ""
            
            if is_template and template_service_instance:
                try:
                    # 문서 ID나 템플릿 이름으로 SQLite에서 템플릿 검색
                    doc_id = doc.get("id")
                    template_name = doc.get("template_name") or doc.get("filename", "").rsplit('.', 1)[0]
                    template_type_str = doc.get("template_type", "")
                    
                    # SQLite에서 템플릿 존재 여부 확인
                    template_exists = False
                    
                    if template_id:
                        # 템플릿 ID로 직접 조회
                        template = await template_service_instance.get_template(template_id)
                        template_exists = template is not None
                    elif template_type_str and template_name:
                        # 템플릿 타입과 이름으로 검색
                        from app.domain.entities.template_entities import TemplateType
                        try:
                            template_type_enum = TemplateType(template_type_str)
                            templates = await template_service_instance.get_templates_by_filters(
                                template_type=template_type_enum,
                                site_id=doc.get("site_id"),
                                limit=10
                            )
                            # 이름이 유사한 템플릿 찾기
                            for template in templates:
                                if template_name in template.name or template.name in template_name:
                                    template_exists = True
                                    doc["template_id"] = template.id  # 템플릿 ID 업데이트
                                    break
                        except ValueError:
                            # 잘못된 템플릿 타입인 경우
                            pass
                    
                    doc["storage_status"]["rdb"] = template_exists
                    
                except Exception as e:
                    logger.warning(f"템플릿 존재 확인 실패 (문서 ID: {doc.get('id')}): {e}")
                    doc["storage_status"]["rdb"] = False
            else:
                doc["storage_status"]["rdb"] = False
            
            # 저장 위치 정보 추가 (아이콘과 함께)
            storage_locations = []
            storage_locations_with_icons = []
            
            if doc["storage_status"]["vector_db"]:
                storage_locations.append("Vector DB (ChromaDB)")
                storage_locations_with_icons.append("🔍 Vector DB")
            if doc["storage_status"]["rdb"]:
                storage_locations.append("SQLite (템플릿)")
                storage_locations_with_icons.append("🗃️ SQLite")
            if doc["storage_status"]["file_storage"]:
                storage_locations.append("File System")
                storage_locations_with_icons.append("📁 File System")
            
            doc["storage_status"]["locations"] = storage_locations
            doc["storage_status"]["locations_with_icons"] = storage_locations_with_icons
            doc["storage_status"]["locations_text"] = " + ".join(storage_locations) if storage_locations else "없음"
            doc["storage_status"]["locations_text_with_icons"] = " + ".join(storage_locations_with_icons) if storage_locations_with_icons else "❌ 없음"
            
            # 저장 상태 요약
            storage_count = sum([
                doc["storage_status"]["vector_db"],
                doc["storage_status"]["rdb"], 
                doc["storage_status"]["file_storage"]
            ])
            doc["storage_status"]["total_storages"] = storage_count
            doc["storage_status"]["is_complete"] = storage_count >= 2  # 2곳 이상 저장되면 완전한 상태
            doc["storage_status"]["is_hybrid"] = storage_count > 1  # 하이브리드 저장 여부
            
            # 저장 상태 레벨 (UI에서 사용할 수 있도록)
            if storage_count == 0:
                doc["storage_status"]["level"] = "none"
                doc["storage_status"]["level_text"] = "저장되지 않음"
                doc["storage_status"]["level_color"] = "red"
                doc["storage_status"]["level_icon"] = "❌"
                doc["storage_status"]["level_description"] = "문서가 어떤 저장소에도 저장되지 않았습니다"
            elif storage_count == 1:
                doc["storage_status"]["level"] = "single"
                doc["storage_status"]["level_text"] = "단일 저장소"
                doc["storage_status"]["level_color"] = "yellow"
                doc["storage_status"]["level_icon"] = "⚠️"
                doc["storage_status"]["level_description"] = "문서가 하나의 저장소에만 저장되어 있습니다"
            elif storage_count == 2:
                doc["storage_status"]["level"] = "hybrid"
                doc["storage_status"]["level_text"] = "하이브리드 저장"
                doc["storage_status"]["level_color"] = "green"
                doc["storage_status"]["level_icon"] = "✅"
                doc["storage_status"]["level_description"] = "문서가 두 개의 저장소에 저장되어 있습니다 (권장)"
            else:
                doc["storage_status"]["level"] = "full"
                doc["storage_status"]["level_text"] = "전체 저장"
                doc["storage_status"]["level_color"] = "blue"
                doc["storage_status"]["level_icon"] = "🎯"
                doc["storage_status"]["level_description"] = "문서가 모든 저장소에 저장되어 있습니다 (최고)"
            
            # 저장 상태별 상세 정보
            storage_details = {
                "vector_db": {
                    "name": "Vector Database",
                    "description": "검색 및 유사도 분석용",
                    "icon": "🔍",
                    "status": "✅" if doc["storage_status"]["vector_db"] else "❌",
                    "enabled": doc["storage_status"]["vector_db"]
                },
                "rdb": {
                    "name": "SQLite Database", 
                    "description": "템플릿 메타데이터 저장용",
                    "icon": "🗃️",
                    "status": "✅" if doc["storage_status"]["rdb"] else "❌",
                    "enabled": doc["storage_status"]["rdb"]
                },
                "file_storage": {
                    "name": "File System",
                    "description": "원본 파일 보관용",
                    "icon": "📁", 
                    "status": "✅" if doc["storage_status"]["file_storage"] else "❌",
                    "enabled": doc["storage_status"]["file_storage"],
                    "path": doc["storage_status"]["file_path"]
                }
            }
            doc["storage_status"]["details"] = storage_details
            
            # 권장사항 제시
            recommendations = []
            if not doc["storage_status"]["vector_db"]:
                recommendations.append("Vector DB에 문서를 추가하여 검색 기능을 활성화하세요")
            if is_template and not doc["storage_status"]["rdb"]:
                recommendations.append("템플릿을 SQLite에 저장하여 템플릿 관리 기능을 활성화하세요")
            if not doc["storage_status"]["file_storage"]:
                recommendations.append("원본 파일을 파일 시스템에 백업하세요")
            
            doc["storage_status"]["recommendations"] = recommendations
        
        return APIResponse(
            success=True,
            message="문서 목록 조회 성공",
            data={
                "documents": documents_list,
                "total_count": len(documents_map),
                "returned_count": len(documents_list),
                "offset": offset,
                "limit": limit,
                "collection_info": {
                    "total_chunks": collection_info.get("total_chunks", 0),
                    "total_sites": collection_info.get("total_sites", 0),
                    "document_types": collection_info.get("document_types", {})
                }
            }
        )
        
    except Exception as e:
        logger.error(f"문서 목록 조회 실패: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"문서 목록 조회 실패: {str(e)}"
        )


@router.get(
    "/{document_id}/content",
    response_model=APIResponse,
    summary="문서 내용 조회",
    description="특정 문서의 전체 내용을 조회합니다. 청크로 나누어 저장된 문서를 재결합하여 반환합니다."
)
async def get_document_content(
    document_id: str,
    service: object = Depends(get_vector_store_service)
) -> APIResponse:
    """
    문서 내용 조회
    
    - **document_id**: 조회할 문서 ID
    
    벡터 데이터베이스에서 해당 문서의 모든 청크를 조회하여 원본 문서 내용을 재구성합니다.
    """
    try:
        logger.info(f"문서 내용 조회 요청: {document_id}")
        
        # ChromaDB에서 해당 문서의 모든 청크 조회
        collection = service._collection
        if collection is None:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="벡터 데이터베이스에 연결할 수 없습니다"
            )
        
        # 문서 ID로 필터링하여 모든 청크 조회
        results = collection.get(
            where={"logical_key": document_id},
            include=["documents", "metadatas", "ids"]
        )
        
        if not results["documents"] or len(results["documents"]) == 0:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"문서를 찾을 수 없습니다: {document_id}"
            )
        
        # 청크들을 순서대로 정렬 (chunk_index 기준)
        chunks_data = []
        for i, (doc_content, metadata, chunk_id) in enumerate(zip(
            results["documents"], 
            results["metadatas"], 
            results["ids"]
        )):
            chunk_index = metadata.get("chunk_index", i)
            chunks_data.append({
                "chunk_index": chunk_index,
                "content": doc_content,
                "metadata": metadata,
                "chunk_id": chunk_id
            })
        
        # chunk_index로 정렬
        chunks_data.sort(key=lambda x: x["chunk_index"])
        
        # 문서 내용 재결합
        full_content = ""
        document_metadata = {}
        
        for chunk in chunks_data:
            full_content += chunk["content"]
            if not document_metadata:  # 첫 번째 청크의 메타데이터 사용
                document_metadata = chunk["metadata"].copy()
        
        # 문서 기본 정보
        document_info = {
            "document_id": document_id,
            "title": document_metadata.get("title", document_metadata.get("original_filename", "제목 없음")),
            "content": full_content,
            "total_chunks": len(chunks_data),
            "content_length": len(full_content),
            "doc_type": document_metadata.get("doc_type", "unknown"),
            "site_id": document_metadata.get("site_id"),
            "created_at": document_metadata.get("created_at"),
            "file_path": document_metadata.get("file_path"),
            "original_filename": document_metadata.get("original_filename"),
            "metadata": document_metadata
        }
        
        logger.info(f"문서 내용 조회 완료: {document_id}, {len(chunks_data)}개 청크, {len(full_content)}자")
        
        return APIResponse(
            success=True,
            message=f"문서 내용 조회 성공: {document_info['title']}",
            data=document_info
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"문서 내용 조회 실패: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"문서 내용 조회 실패: {str(e)}"
        )


@router.get(
    "/stats",
    response_model=APIResponse,
    summary="문서 통계",
    description="벡터 데이터베이스의 문서 통계를 조회합니다."
)
async def get_document_stats(
    service: object = Depends(get_vector_store_service)
) -> APIResponse:
    """문서 통계 조회"""
    try:
        stats = await service.get_collection_info()
        
        return APIResponse(
            success=True,
            message="문서 통계 조회 성공",
            data=stats
        )
        
    except Exception as e:
        logger.error(f"문서 통계 조회 실패: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"문서 통계 조회 실패: {str(e)}"
        )


@router.patch(
    "/{document_id}/template-type",
    response_model=APIResponse,
    summary="문서 템플릿 타입 업데이트",
    description="기존 문서의 템플릿 타입을 업데이트합니다."
)
async def update_document_template_type(
    document_id: str,
    template_type: str = Form(...),
    template_name: str = Form(None),
    template_description: str = Form(None),
    service: object = Depends(get_vector_store_service)
) -> APIResponse:
    """기존 문서의 템플릿 타입 업데이트"""
    try:
        if not service._collection:
            await service.initialize()
        
        # 문서 존재 확인 및 현재 메타데이터 조회
        existing_results = service._collection.get(
            where={"document_id": document_id},
            include=["metadatas", "documents"]
        )
        
        if not existing_results["ids"]:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"문서를 찾을 수 없습니다: {document_id}"
            )
        
        # 모든 청크에 대해 메타데이터 업데이트
        updated_count = 0
        for i, chunk_id in enumerate(existing_results["ids"]):
            current_metadata = existing_results["metadatas"][i]
            
            # 새로운 메타데이터 구성
            updated_metadata = {
                **current_metadata,
                "template_type": template_type,
                "template_name": template_name if template_name else current_metadata.get("template_name", ""),
                "template_description": template_description if template_description else current_metadata.get("template_description", ""),
                "is_template": True,
                "updated_at": datetime.now().isoformat()
            }
            
            # ChromaDB에서 메타데이터 업데이트
            service._collection.update(
                ids=[chunk_id],
                metadatas=[updated_metadata]
            )
            updated_count += 1
        
        logger.info(f"문서 {document_id}의 템플릿 타입을 {template_type}로 업데이트 완료 ({updated_count}개 청크)")
        
        return APIResponse(
            success=True,
            message=f"문서 템플릿 타입 업데이트 완료",
            data={
                "document_id": document_id,
                "template_type": template_type,
                "template_name": template_name,
                "template_description": template_description,
                "updated_chunks": updated_count
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"문서 템플릿 타입 업데이트 실패: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"문서 템플릿 타입 업데이트 실패: {str(e)}"
        )


@router.post(
    "/batch-update-template-types",
    response_model=APIResponse,
    summary="일괄 템플릿 타입 업데이트",
    description="여러 문서의 템플릿 타입을 일괄 업데이트합니다."
)
async def batch_update_template_types(
    updates: List[Dict[str, str]],  # [{"document_id": "...", "template_type": "...", "template_name": "...", "template_description": "..."}]
    service: object = Depends(get_vector_store_service)
) -> APIResponse:
    """여러 문서의 템플릿 타입을 일괄 업데이트"""
    try:
        if not service._collection:
            await service.initialize()
        
        results = []
        
        for update in updates:
            document_id = update.get("document_id")
            template_type = update.get("template_type")
            template_name = update.get("template_name", "")
            template_description = update.get("template_description", "")
            
            if not document_id or not template_type:
                results.append({
                    "document_id": document_id,
                    "status": "failed",
                    "error": "document_id와 template_type은 필수입니다"
                })
                continue
            
            try:
                # 문서 존재 확인 및 현재 메타데이터 조회
                existing_results = service._collection.get(
                    where={"document_id": document_id},
                    include=["metadatas"]
                )
                
                if not existing_results["ids"]:
                    results.append({
                        "document_id": document_id,
                        "status": "failed",
                        "error": "문서를 찾을 수 없습니다"
                    })
                    continue
                
                # 모든 청크에 대해 메타데이터 업데이트
                updated_count = 0
                for i, chunk_id in enumerate(existing_results["ids"]):
                    current_metadata = existing_results["metadatas"][i]
                    
                    # 새로운 메타데이터 구성
                    updated_metadata = {
                        **current_metadata,
                        "template_type": template_type,
                        "template_name": template_name,
                        "template_description": template_description,
                        "is_template": True,
                        "updated_at": datetime.now().isoformat()
                    }
                    
                    # ChromaDB에서 메타데이터 업데이트
                    service._collection.update(
                        ids=[chunk_id],
                        metadatas=[updated_metadata]
                    )
                    updated_count += 1
                
                results.append({
                    "document_id": document_id,
                    "status": "success",
                    "template_type": template_type,
                    "updated_chunks": updated_count
                })
                
            except Exception as e:
                results.append({
                    "document_id": document_id,
                    "status": "failed",
                    "error": str(e)
                })
        
        success_count = len([r for r in results if r["status"] == "success"])
        failed_count = len([r for r in results if r["status"] == "failed"])
        
        return APIResponse(
            success=True,
            message=f"일괄 업데이트 완료 (성공: {success_count}, 실패: {failed_count})",
            data={
                "total": len(updates),
                "success_count": success_count,
                "failed_count": failed_count,
                "results": results
            }
        )
        
    except Exception as e:
        logger.error(f"일괄 템플릿 타입 업데이트 실패: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"일괄 템플릿 타입 업데이트 실패: {str(e)}"
        )


@router.get(
    "/{document_id}/storage-status",
    response_model=APIResponse,
    summary="문서 저장 상태 상세 조회",
    description="특정 문서의 하이브리드 저장 상태를 상세하게 조회합니다."
)
async def get_document_storage_status(
    document_id: str,
    service: object = Depends(get_vector_store_service)
) -> APIResponse:
    """특정 문서의 저장 상태 상세 조회"""
    try:
        if not service._collection:
            await service.initialize()
        
        # Vector DB에서 문서 조회
        results = service._collection.get(
            where={"document_id": document_id},
            include=["metadatas", "documents"]
        )
        
        if not results["ids"]:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"문서를 찾을 수 없습니다: {document_id}"
            )
        
        # 첫 번째 청크의 메타데이터 사용 (문서 정보는 모든 청크에 동일)
        metadata = results["metadatas"][0]
        
        # 문서 정보 구성
        doc_info = {
            "id": document_id,
            "title": metadata.get("title", "제목 없음"),
            "filename": metadata.get("original_filename", metadata.get("title", "파일명 없음")),
            "type": metadata.get("doc_type", "unknown"),
            "site_id": metadata.get("site_id"),
            "created_at": metadata.get("created_at"),
            "template_type": metadata.get("template_type", ""),
            "template_name": metadata.get("template_name", ""),
            "is_template": metadata.get("is_template", False),
            "total_chunks": len(results["ids"])
        }
        
        # 저장 상태 확인 (위에서 구현한 로직과 동일)
        template_service_instance = get_template_service()
        
        # Vector DB 상태 (이미 조회했으므로 True)
        storage_status = {"vector_db": True}
        
        # 파일 저장 상태 확인
        file_path = metadata.get("file_path")
        file_exists = False
        
        if file_path and os.path.exists(file_path):
            file_exists = True
        else:
            # 업로드 디렉토리에서 패턴 매칭으로 파일 찾기
            for ext in ['.md', '.txt', '.pdf', '.doc', '.docx', '.html']:
                pattern_path = f"{UPLOAD_DIRECTORY}/{document_id}{ext}"
                if os.path.exists(pattern_path):
                    file_path = pattern_path
                    file_exists = True
                    break
        
        storage_status["file_storage"] = file_exists
        storage_status["file_path"] = file_path if file_exists else None
        
        # RDB 저장 상태 확인
        is_template = doc_info.get("is_template", False) or doc_info.get("template_type", "") != ""
        template_exists = False
        
        if is_template and template_service_instance:
            try:
                template_type_str = doc_info.get("template_type", "")
                template_name = doc_info.get("template_name") or doc_info.get("filename", "").rsplit('.', 1)[0]
                
                if template_type_str and template_name:
                    from app.domain.entities.template_entities import TemplateType
                    try:
                        template_type_enum = TemplateType(template_type_str)
                        templates = await template_service_instance.get_templates_by_filters(
                            template_type=template_type_enum,
                            site_id=doc_info.get("site_id"),
                            limit=10
                        )
                        for template in templates:
                            if template_name in template.name or template.name in template_name:
                                template_exists = True
                                break
                    except ValueError:
                        pass
            except Exception as e:
                logger.warning(f"템플릿 존재 확인 실패: {e}")
        
        storage_status["rdb"] = template_exists
        
        # 저장 상태 요약 정보 생성 (위에서 구현한 로직 재사용)
        storage_count = sum([
            storage_status["vector_db"],
            storage_status["rdb"],
            storage_status["file_storage"]
        ])
        
        # 상태 레벨 결정
        if storage_count == 0:
            level = "none"
            level_text = "저장되지 않음"
            level_color = "red"
            level_icon = "❌"
        elif storage_count == 1:
            level = "single"
            level_text = "단일 저장소"
            level_color = "yellow"
            level_icon = "⚠️"
        elif storage_count == 2:
            level = "hybrid"
            level_text = "하이브리드 저장"
            level_color = "green"
            level_icon = "✅"
        else:
            level = "full"
            level_text = "전체 저장"
            level_color = "blue"
            level_icon = "🎯"
        
        # 저장 위치 정보
        locations = []
        if storage_status["vector_db"]:
            locations.append("🔍 Vector DB")
        if storage_status["rdb"]:
            locations.append("🗃️ SQLite")
        if storage_status["file_storage"]:
            locations.append("📁 File System")
        
        return APIResponse(
            success=True,
            message="문서 저장 상태 조회 성공",
            data={
                "document": doc_info,
                "storage_status": {
                    **storage_status,
                    "total_storages": storage_count,
                    "is_hybrid": storage_count > 1,
                    "is_complete": storage_count >= 2,
                    "level": level,
                    "level_text": level_text,
                    "level_color": level_color,
                    "level_icon": level_icon,
                    "locations": locations,
                    "locations_text": " + ".join(locations) if locations else "❌ 없음"
                }
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"문서 저장 상태 조회 실패: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"문서 저장 상태 조회 실패: {str(e)}"
        ) 