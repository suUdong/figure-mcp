"""
Word 문서 프로세서
"""

from typing import Dict, Any, List
from docx import Document
from docx.table import Table
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
import re

from .base import BaseProcessor


class WordProcessor(BaseProcessor):
    """Word 문서 처리기"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.docx', '.doc']
        self.supported_mime_types = [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword'
        ]
    
    async def extract_text(self, file_path: str) -> str:
        """Word 문서에서 텍스트 추출"""
        try:
            doc = Document(file_path)
            text_parts = []
            
            # 문서의 모든 요소를 순서대로 처리
            for element in doc.element.body:
                if isinstance(element, CT_P):
                    # 단락 처리
                    paragraph = next(p for p in doc.paragraphs if p.element == element)
                    if paragraph.text.strip():
                        # 스타일 정보 포함
                        style_name = paragraph.style.name if paragraph.style else None
                        if style_name and 'Heading' in style_name:
                            text_parts.append(f"\n## {paragraph.text.strip()}\n")
                        else:
                            text_parts.append(paragraph.text.strip())
                            
                elif isinstance(element, CT_Tbl):
                    # 테이블 처리
                    table = next(t for t in doc.tables if t.element == element)
                    table_text = self._extract_table_text(table)
                    if table_text:
                        text_parts.append(f"\n[표]\n{table_text}\n")
            
            return self.clean_text("\n".join(text_parts))
            
        except Exception as e:
            raise ValueError(f"Word 문서 텍스트 추출 실패: {str(e)}")
    
    async def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Word 문서에서 메타데이터 추출"""
        try:
            doc = Document(file_path)
            core_props = doc.core_properties
            
            metadata = {
                'processor_type': 'word',
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'section_count': len(doc.sections),
                'styles_used': [],
                'headings': [],
                'has_images': False,
                'word_count': 0,
            }
            
            # 문서 속성
            if core_props.title:
                metadata['title'] = core_props.title
            if core_props.author:
                metadata['author'] = core_props.author
            if core_props.subject:
                metadata['subject'] = core_props.subject
            if core_props.created:
                metadata['created'] = core_props.created.isoformat()
            if core_props.modified:
                metadata['modified'] = core_props.modified.isoformat()
            
            # 스타일 및 헤딩 분석
            styles_used = set()
            headings = []
            total_words = 0
            
            for paragraph in doc.paragraphs:
                if paragraph.style:
                    styles_used.add(paragraph.style.name)
                    
                # 헤딩 추출
                if paragraph.style and 'Heading' in paragraph.style.name:
                    headings.append({
                        'level': paragraph.style.name,
                        'text': paragraph.text.strip()
                    })
                
                # 단어 수 계산
                words = len(paragraph.text.split())
                total_words += words
            
            metadata['styles_used'] = list(styles_used)
            metadata['headings'] = headings
            metadata['word_count'] = total_words
            
            # 이미지 확인 (간단한 방법)
            for rel in doc.part.rels.values():
                if 'image' in rel.target_ref:
                    metadata['has_images'] = True
                    break
            
            return metadata
            
        except Exception as e:
            raise ValueError(f"Word 메타데이터 추출 실패: {str(e)}")
    
    def _extract_table_text(self, table: Table) -> str:
        """테이블을 텍스트로 변환"""
        if not table.rows:
            return ""
        
        text_parts = []
        
        # 첫 번째 행을 헤더로 간주
        if table.rows:
            header_row = table.rows[0]
            headers = []
            for cell in header_row.cells:
                headers.append(cell.text.strip())
            
            if any(headers):  # 헤더가 있는 경우
                text_parts.append(" | ".join(headers))
                text_parts.append("-" * min(80, len(" | ".join(headers))))
        
        # 데이터 행들
        for row in table.rows[1:]:  # 첫 번째 행 제외
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            
            if any(row_data):  # 빈 행이 아닌 경우
                text_parts.append(" | ".join(row_data))
        
        return "\n".join(text_parts)
    
    async def extract_structured_data(self, file_path: str) -> List[Dict[str, Any]]:
        """Word에서 구조화된 데이터 추출"""
        try:
            doc = Document(file_path)
            structured_data = []
            
            # 헤딩 구조 추출
            headings_structure = self._extract_headings_structure(doc)
            if headings_structure:
                structured_data.append({
                    'type': 'outline',
                    'title': '문서 구조',
                    'data': headings_structure
                })
            
            # 테이블 구조 추출
            for i, table in enumerate(doc.tables):
                table_data = self._extract_table_structure(table)
                if table_data:
                    table_data['table_index'] = i
                    structured_data.append(table_data)
            
            # 목록 구조 추출
            lists = self._extract_lists(doc)
            for i, list_data in enumerate(lists):
                list_data['list_index'] = i
                structured_data.append(list_data)
            
            return structured_data
            
        except Exception as e:
            raise ValueError(f"Word 구조화된 데이터 추출 실패: {str(e)}")
    
    def _extract_headings_structure(self, doc: Document) -> List[Dict[str, Any]]:
        """헤딩 구조 추출"""
        headings = []
        
        for paragraph in doc.paragraphs:
            if paragraph.style and 'Heading' in paragraph.style.name:
                level_match = re.search(r'Heading (\d+)', paragraph.style.name)
                level = int(level_match.group(1)) if level_match else 1
                
                headings.append({
                    'level': level,
                    'text': paragraph.text.strip(),
                    'style': paragraph.style.name
                })
        
        return headings
    
    def _extract_table_structure(self, table: Table) -> Dict[str, Any]:
        """테이블 구조 추출"""
        if not table.rows:
            return None
        
        table_data = {
            'type': 'table',
            'rows': len(table.rows),
            'columns': len(table.rows[0].cells) if table.rows else 0,
            'data': []
        }
        
        # 헤더 추출
        if table.rows:
            header_row = table.rows[0]
            headers = [cell.text.strip() for cell in header_row.cells]
            table_data['headers'] = headers
        
        # 데이터 행 추출 (최대 50행)
        max_rows = min(50, len(table.rows))
        for row in table.rows[:max_rows]:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data['data'].append(row_data)
        
        return table_data
    
    def _extract_lists(self, doc: Document) -> List[Dict[str, Any]]:
        """목록 구조 추출"""
        lists = []
        current_list = None
        
        for paragraph in doc.paragraphs:
            # 번호 매김 또는 불릿 포인트 확인
            if self._is_list_item(paragraph):
                if current_list is None:
                    current_list = {
                        'type': 'list',
                        'items': []
                    }
                
                current_list['items'].append({
                    'text': paragraph.text.strip(),
                    'level': self._get_list_level(paragraph)
                })
            else:
                if current_list and current_list['items']:
                    lists.append(current_list)
                    current_list = None
        
        # 마지막 목록 추가
        if current_list and current_list['items']:
            lists.append(current_list)
        
        return lists
    
    def _is_list_item(self, paragraph) -> bool:
        """문단이 목록 항목인지 확인"""
        text = paragraph.text.strip()
        if not text:
            return False
        
        # 번호 매김 확인 (1., 2., a., b., i., ii. 등)
        if re.match(r'^\d+\.', text) or re.match(r'^[a-zA-Z]\.', text) or re.match(r'^[ivx]+\.', text):
            return True
        
        # 불릿 포인트 확인
        if text.startswith('•') or text.startswith('-') or text.startswith('*'):
            return True
        
        # Word의 목록 스타일 확인
        if paragraph.style and 'List' in paragraph.style.name:
            return True
        
        return False
    
    def _get_list_level(self, paragraph) -> int:
        """목록 항목의 레벨 확인"""
        # 들여쓰기나 스타일을 기반으로 레벨 결정
        if paragraph.style and 'List' in paragraph.style.name:
            level_match = re.search(r'List (\d+)', paragraph.style.name)
            if level_match:
                return int(level_match.group(1))
        
        # 기본값
        return 1 